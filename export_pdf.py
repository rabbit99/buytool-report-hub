"""
將 reports/ 下所有 .md 檔批次轉為 HTML（可直接瀏覽器列印 PDF）
以及 .docx（可直接上傳 Google Drive → 開啟為 Google 文件）

用法：
    python export_pdf.py          # 產生 HTML + DOCX
    python export_pdf.py --html   # 只產生 HTML
    python export_pdf.py --docx   # 只產生 DOCX
"""
import argparse
import pathlib
import markdown

REPORTS_DIR = pathlib.Path(__file__).parent / "reports"
CSS = """
body { font-family: "Microsoft JhengHei", "Noto Sans TC", sans-serif; font-size: 14px; line-height: 1.8; max-width: 900px; margin: 40px auto; padding: 0 20px; color: #333; }
h1 { font-size: 22px; border-bottom: 2px solid #333; padding-bottom: 6px; }
h2 { font-size: 18px; margin-top: 28px; color: #1a5276; }
h3 { font-size: 15px; color: #2e4053; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; }
th, td { border: 1px solid #bbb; padding: 6px 10px; text-align: left; }
th { background: #eaf2f8; }
code { background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-size: 13px; }
pre { background: #f4f4f4; padding: 14px; border-radius: 5px; overflow-x: auto; white-space: pre-wrap; }
blockquote { border-left: 4px solid #aed6f1; padding-left: 14px; color: #555; margin: 14px 0; background: #fafafa; }
@media print { body { margin: 20px; } }
"""


def md_to_html(md_path: pathlib.Path) -> str:
    text = md_path.read_text(encoding="utf-8")
    html_body = markdown.markdown(text, extensions=["tables", "fenced_code"])
    return (
        "<!DOCTYPE html>\n"
        f"<html><head><meta charset='utf-8'><title>{md_path.stem}</title>"
        f"<style>{CSS}</style></head><body>{html_body}</body></html>"
    )


def export_html(md_files: list[pathlib.Path]):
    for md_path in md_files:
        html = md_to_html(md_path)
        out = md_path.with_suffix(".html")
        out.write_text(html, encoding="utf-8")
        print(f"✅ {out.name}")


def export_docx(md_files: list[pathlib.Path]):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("需要安裝 python-docx：pip install python-docx")
        return

    for md_path in md_files:
        text = md_path.read_text(encoding="utf-8")
        doc = Document()

        # 設定預設字型
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft JhengHei"
        font.size = Pt(11)

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            if stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("> "):
                p = doc.add_paragraph(stripped[2:])
                p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else p.style
            elif stripped.startswith("|") and stripped.endswith("|"):
                # 簡化處理：表格行直接當段落
                doc.add_paragraph(stripped, style="Normal")
            else:
                doc.add_paragraph(stripped)

        out = md_path.with_suffix(".docx")
        doc.save(str(out))
        print(f"✅ {out.name}")


def main():
    from vendor_config import list_vendors

    parser = argparse.ArgumentParser(description="匯出 MD 報告為 HTML / DOCX")
    parser.add_argument("--html", action="store_true", help="只匯出 HTML")
    parser.add_argument("--docx", action="store_true", help="只匯出 DOCX")
    parser.add_argument(
        "--vendor", "-v", type=str, default=None,
        help=f"指定廠商（{', '.join(list_vendors())}），僅匯出該廠商的報告",
    )
    args = parser.parse_args()

    if args.vendor:
        report_dir = REPORTS_DIR / args.vendor
        if not report_dir.exists():
            print(f"reports/{args.vendor}/ 資料夾不存在")
            return
        md_files = sorted(report_dir.glob("*.md"))
        label = f"reports/{args.vendor}/"
    else:
        # 搜尋所有子資料夾
        md_files = sorted(REPORTS_DIR.rglob("*.md"))
        label = "reports/"

    if not md_files:
        print(f"{label} 沒有 .md 檔案")
        return

    do_both = not args.html and not args.docx

    if do_both or args.html:
        print("── HTML ──")
        export_html(md_files)

    if do_both or args.docx:
        print("── DOCX ──")
        export_docx(md_files)

    if do_both or args.html:
        print("\n💡 HTML → PDF：用瀏覽器開啟 .html → Ctrl+P → 儲存為 PDF")
    if do_both or args.docx:
        print("💡 DOCX → Google 文件：上傳 .docx 到 Google Drive → 開啟為 Google 文件")


if __name__ == "__main__":
    main()
